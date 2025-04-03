import streamlit as st
import datetime
import os
from google import genai
from docx import Document
from docx.shared import Inches
from io import BytesIO
import pandas as pd

# --- Configure your API key securely ---
API_KEY = os.getenv("GOOGLE_API_KEY", "AIzaSyBHuqIqb_udt5qKlP0fPCwqMu2JZ4vQVV4")
client = genai.Client(api_key=API_KEY)

# --- Analytics Setup ---
if "usage_stats" not in st.session_state:
    st.session_state.usage_stats = []

# --- Set page config and logo ---
st.set_page_config(page_title="Legal Document Drafter", page_icon="📜")
st.title("📜 AI Legal Document Drafter")

# --- Sidebar Navigation ---
page = st.sidebar.radio("Navigate", ["📄 Draft Document", "📊 Analytics Dashboard"])

# --- Define templates ---
DOC_TEMPLATES = {
    "NDA": ["Disclosing Party", "Receiving Party", "Jurisdiction", "Start Date", "End Date"],
    "Power of Attorney": ["Grantor", "Agent", "Scope of Authority", "Start Date", "End Date"],
    "Legal Dispute Complaint": ["Plaintiff", "Defendant", "Jurisdiction", "Incident Date", "Description", "Damages", "Monetary Relief"],
    "Lease Agreement": ["Lessor", "Lessee", "Property Address", "Rent Amount", "Lease Duration", "Terms"],
    "Internship Contract": ["Intern Name", "Company", "Duration", "Stipend", "Scope of Work"],
    "Consulting Agreement": ["Client", "Consultant", "Fee", "Deliverables", "Termination Clause"],
    "Will/Testament": ["Testator", "Beneficiaries", "Assets", "Executor", "Witness Information"]
}

if page == "📄 Draft Document":
    # --- Select document type ---
    doc_type = st.selectbox("Choose a Legal Document Type:", list(DOC_TEMPLATES.keys()))

    # --- Input form ---
    with st.form("doc_form"):
        st.subheader(f"Enter details for {doc_type}")
        user_inputs = {}
        for field in DOC_TEMPLATES[doc_type]:
            if "Date" in field:
                user_inputs[field] = st.date_input(field, value=datetime.date.today())
            elif any(keyword in field for keyword in ["Description", "Scope", "Damages", "Terms", "Deliverables", "Assets", "Witness"]):
                user_inputs[field] = st.text_area(field)
            else:
                user_inputs[field] = st.text_input(field)
        submitted = st.form_submit_button("Generate Document")

    # --- Generate document ---
    if submitted:
        prompt = f"Draft a formal {doc_type} with the following details:\n"
        for key, val in user_inputs.items():
            prompt += f"{key}: {val}\n"

        with st.spinner("Generating legal document..."):
            try:
                response = client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=prompt
                )
                draft = response.text
                st.success("✅ Document Generated!")
                st.text_area("Generated Document", draft, height=400)

                # Create Word Document
                doc = Document()

                doc.add_heading(f"{doc_type} Draft", level=1)
                for line in draft.split("\n"):
                    doc.add_paragraph(line)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                filename = f"{doc_type.replace(' ', '_').lower()}_draft.docx"
                st.download_button(
                    label="📥 Download as Word Document",
                    data=buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Track usage
                st.session_state.usage_stats.append({"Document": doc_type, "Time": datetime.datetime.now()})

            except Exception as e:
                import traceback
                st.error("❌ An error occurred while generating the document:")
                st.code(traceback.format_exc())

elif page == "📊 Analytics Dashboard":
    st.header("📊 Document Generation Analytics")

    if st.session_state.usage_stats:
        df = pd.DataFrame(st.session_state.usage_stats)
        doc_counts = df["Document"].value_counts()

        st.bar_chart(doc_counts)
        st.metric("📄 Total Documents Generated", len(df))
    else:
        st.info("No documents generated yet in this session.")